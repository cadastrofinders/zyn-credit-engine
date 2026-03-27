"""
ZYN Capital — Gerador de MAC v4.1 (.docx)
Memorando de Análise de Crédito com identidade visual ZYN — World-Class Edition.
Template unificado baseado em 4 MACs reais:
  - Wolney/Plantae (NC/CCB), Ivanoff (SLB Agro),
  - Baron55 (CRI Imobiliário), Saul Francisco (SLB PDF)

Dependência: python-docx >= 1.1
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.table import _Cell

# ---------------------------------------------------------------------------
# Color Palette — ZYN Capital v4.1
# ---------------------------------------------------------------------------
NAVY_HEX = "223040"
DARK_SLATE_HEX = "3A4F63"
MUTED_GRAY_HEX = "8B9197"
LIGHT_LABEL_HEX = "A8B4C0"
LIGHT_GRAY_BG_HEX = "F2F4F6"
LIGHT_GRAY_BG2_HEX = "E8ECF0"
ALT_ROW_HEX = "F7F8FA"
BORDER_GRAY_HEX = "D0D5DA"
BORDER_LIGHT_HEX = "E0E4E8"
SEPARATOR_HEX = "D0D5DA"
GREEN_BG_HEX = "EAF4EE"
GREEN_SOLID_HEX = "2E7D4F"
DARK_GREEN_HEX = "1E6B42"
YELLOW_BG_HEX = "FEF9E7"
YELLOW_BORDER_HEX = "E6A817"
DARK_GOLD_HEX = "7D6608"
RED_BG_HEX = "FDECEA"
RED_LIGHT_BG_HEX = "FBE4E4"
DARK_RED_HEX = "922B21"
RED_SOLID_HEX = "C0392B"
DARK_BLUE_HEX = "1A5276"
BLUE_BG_HEX = "E8F0FE"
WHITE_HEX = "FFFFFF"
CREAM_HEX = "FFF8E7"

NAVY = RGBColor(34, 48, 64)
DARK_SLATE = RGBColor(58, 79, 99)
MUTED_GRAY = RGBColor(139, 145, 151)
LIGHT_LABEL = RGBColor(168, 180, 192)
DARK_GREEN = RGBColor(30, 107, 66)
DARK_GOLD = RGBColor(125, 102, 8)
DARK_RED = RGBColor(146, 43, 33)
DARK_BLUE = RGBColor(26, 82, 118)
WHITE = RGBColor(255, 255, 255)

# KPI strip column colors — 5 distinct backgrounds
KPI_COLORS = [NAVY_HEX, DARK_SLATE_HEX, "2E7D4F", "5B4A1E", "6B2D1F"]

# ---------------------------------------------------------------------------
# Section titles — Unified MAC v4 (13 sections)
# ---------------------------------------------------------------------------
SECTION_TITLES = {
    "tomador": "I. TOMADOR",
    "patrimonio": "II. PATRIMÔNIO",
    "producao": "III. PRODUÇÃO E ATIVIDADE OPERACIONAL",
    "capital": "IV. CAPITAL E ESTRUTURA DE ENDIVIDAMENTO",
    "operacao": "V. OPERAÇÃO",
    "estrutura": "VI. ESTRUTURA DA OPERAÇÃO",
    "pagamento": "VII. MODELAGEM FINANCEIRA — CRONOGRAMA",
    "capacidade": "VIII. ANÁLISE FINANCEIRA — CAPACIDADE DE PAGAMENTO",
    "impacto": "IX. IMPACTO FINANCEIRO",
    "onus": "X. MAPA DE ÔNUS E GARANTIAS",
    "riscos": "XI. RISCOS, MITIGAÇÕES E COVENANTS",
    "cronograma": "XII. CRONOGRAMA DE EXECUÇÃO",
    "conclusao": "XIII. CONCLUSÃO E RECOMENDAÇÃO — ZYN CAPITAL",
}

# Alert type configs: (fill_hex, border_hex, title_color_hex)
ALERT_STYLES = {
    "warning": (CREAM_HEX, YELLOW_BORDER_HEX, DARK_GOLD_HEX),
    "info": (BLUE_BG_HEX, DARK_BLUE_HEX, DARK_BLUE_HEX),
    "risk": (RED_LIGHT_BG_HEX, RED_SOLID_HEX, DARK_RED_HEX),
    "positive": (GREEN_BG_HEX, GREEN_SOLID_HEX, DARK_GREEN_HEX),
    "critical": (RED_LIGHT_BG_HEX, RED_SOLID_HEX, DARK_RED_HEX),
}

# Risk severity styles: (text_hex, fill_hex) — v4.1 uses solid fills for high severity
RISK_LEVEL_STYLES = {
    "CRÍTICO": (WHITE_HEX, NAVY_HEX),
    "ALTO": (WHITE_HEX, RED_SOLID_HEX),
    "MÉDIO": (DARK_GOLD_HEX, YELLOW_BG_HEX),
    "BAIXO": (WHITE_HEX, GREEN_SOLID_HEX),
}

# Rating scale AAA-D → alert styles
RATING_ALERT_MAP = {
    "AAA": "positive", "AA": "positive", "A": "positive",
    "BBB": "info", "BB": "warning",
    "B": "warning", "C": "risk", "D": "risk",
}

# Rating colors for the large badge
RATING_BADGE_COLORS = {
    "AAA": (WHITE_HEX, GREEN_SOLID_HEX), "AA": (WHITE_HEX, GREEN_SOLID_HEX),
    "A": (WHITE_HEX, GREEN_SOLID_HEX),
    "BBB": (WHITE_HEX, DARK_BLUE_HEX), "BB": (DARK_GOLD_HEX, YELLOW_BG_HEX),
    "B": (DARK_GOLD_HEX, YELLOW_BG_HEX),
    "C": (WHITE_HEX, RED_SOLID_HEX), "D": (WHITE_HEX, DARK_RED_HEX),
}

# Status badge styles
STATUS_STYLES = {
    "SAUDÁVEL": (DARK_GREEN_HEX, GREEN_BG_HEX),
    "COBERTO": (DARK_GREEN_HEX, GREEN_BG_HEX),
    "EXCELENTE": (DARK_GREEN_HEX, GREEN_BG_HEX),
    "ATENÇÃO": (DARK_GOLD_HEX, YELLOW_BG_HEX),
    "MONITORAR": (DARK_GOLD_HEX, YELLOW_BG_HEX),
    "CRÍTICO": (DARK_RED_HEX, RED_BG_HEX),
}

# Table total width in dxa
TABLE_WIDTH_DXA = 9746


# ═══════════════════════════════════════════════════════════════════════════
# Helper Functions
# ═══════════════════════════════════════════════════════════════════════════

def _hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_shading(cell: _Cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(
    cell: _Cell,
    top: Optional[Tuple[str, int, str]] = None,
    bottom: Optional[Tuple[str, int, str]] = None,
    left: Optional[Tuple[str, int, str]] = None,
    right: Optional[Tuple[str, int, str]] = None,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if spec is not None:
            val, sz, color = spec
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
    tc_pr.append(borders)


def _set_cell_no_borders(cell: _Cell) -> None:
    _set_cell_borders(
        cell,
        top=("none", 0, WHITE_HEX),
        bottom=("none", 0, WHITE_HEX),
        left=("none", 0, WHITE_HEX),
        right=("none", 0, WHITE_HEX),
    )


def _set_cell_margins(cell: _Cell, top: int, left: int, bottom: int, right: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("start", left), ("bottom", bottom), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tc_pr.append(margins)


def _set_run_font(
    run, size_pt: float, color_hex: str, bold: bool = False, italic: bool = False,
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    r, g, b = _hex_to_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)


def _set_table_width(table, width_dxa: int = TABLE_WIDTH_DXA) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(width_dxa))
    tw.set(qn("w:type"), "dxa")
    for existing in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(existing)
    tbl_pr.append(tw)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), WHITE_HEX)
        borders.append(el)
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


def _set_col_width(cell: _Cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tw = OxmlElement("w:tcW")
    tw.set(qn("w:w"), str(width_dxa))
    tw.set(qn("w:type"), "dxa")
    for existing in tc_pr.findall(qn("w:tcW")):
        tc_pr.remove(existing)
    tc_pr.append(tw)


def _clear_cell(cell: _Cell) -> None:
    for p in cell.paragraphs:
        p_element = p._element
        p_element.getparent().remove(p_element)


def _add_para_to_cell(cell: _Cell, text: str = "", **run_kwargs):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if text:
        run = p.add_run(text)
        if run_kwargs:
            _set_run_font(run, **run_kwargs)
    return p


def _add_horizontal_rule(doc, color_hex: str = SEPARATOR_HEX, thickness: int = 4) -> None:
    """Add a thin horizontal line between sections using a 1-cell table."""
    table = doc.add_table(rows=1, cols=1)
    _set_table_width(table)
    _remove_table_borders(table)
    cell = table.cell(0, 0)
    _set_cell_margins(cell, top=0, left=0, bottom=0, right=0)
    _set_col_width(cell, TABLE_WIDTH_DXA)
    _set_cell_borders(cell,
                      top=("single", thickness, color_hex),
                      bottom=("none", 0, WHITE_HEX),
                      left=("none", 0, WHITE_HEX),
                      right=("none", 0, WHITE_HEX))
    _clear_cell(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(1)


def _add_spacer(doc, pts: float = 6) -> None:
    """Add vertical spacing between elements."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(1)


# ═══════════════════════════════════════════════════════════════════════════
# Formatting helpers
# ═══════════════════════════════════════════════════════════════════════════

def _format_currency(value) -> str:
    if value is None:
        return "N/D"
    try:
        v = float(value)
        formatted = f"{v:,.0f}".replace(",", ".")
        return f"R$ {formatted}"
    except (ValueError, TypeError):
        return str(value)


def _format_pct(value) -> str:
    if value is None:
        return "N/D"
    try:
        v = float(value)
        formatted = f"{v:.1f}".replace(".", ",")
        return f"{formatted}%"
    except (ValueError, TypeError):
        return str(value)


def _safe(val, default: str = "N/D") -> str:
    if val is None or val == "":
        return default
    return str(val)


# ═══════════════════════════════════════════════════════════════════════════
# Document Component Builders — v4.1 Enhanced Visual Quality
# ═══════════════════════════════════════════════════════════════════════════

def _add_header_bar(doc, parametros: dict) -> None:
    """Branded header: ZYN CAPITAL | DISCLAIMER | Data — 3-column table with navy bottom border."""
    table = doc.add_table(rows=1, cols=3)
    _set_table_width(table)
    _remove_table_borders(table)

    widths = [2400, 5146, 2200]
    tipo_op = _safe(parametros.get("tipo_operacao"), "MEMORANDO")
    tomador = _safe(parametros.get("tomador"), "")
    data_str = datetime.now().strftime("%B / %Y").capitalize()

    # Col 0 — ZYN CAPITAL (branded)
    c0 = table.cell(0, 0)
    _set_cell_no_borders(c0)
    _set_col_width(c0, widths[0])
    _set_cell_margins(c0, top=80, left=80, bottom=80, right=40)
    _clear_cell(c0)
    p0 = c0.add_paragraph()
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0)
    run_zyn = p0.add_run("ZYN")
    _set_run_font(run_zyn, size_pt=13, color_hex=NAVY_HEX, bold=True)
    run_cap = p0.add_run(" CAPITAL")
    _set_run_font(run_cap, size_pt=13, color_hex=MUTED_GRAY_HEX, bold=False)

    # Col 1 — Disclaimer
    c1 = table.cell(0, 1)
    _set_cell_no_borders(c1)
    _set_col_width(c1, widths[1])
    _set_cell_margins(c1, top=80, left=40, bottom=80, right=40)
    _clear_cell(c1)
    p1 = c1.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    disclaimer = f"DISCLAIMER E TERMOS DE CONFIDENCIALIDADE — {tipo_op} | {tomador}".upper()
    run_disc = p1.add_run(disclaimer)
    _set_run_font(run_disc, size_pt=6, color_hex=MUTED_GRAY_HEX)

    # Col 2 — Data
    c2 = table.cell(0, 2)
    _set_cell_no_borders(c2)
    _set_col_width(c2, widths[2])
    _set_cell_margins(c2, top=80, left=40, bottom=80, right=80)
    _clear_cell(c2)
    p2 = c2.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run_data = p2.add_run(data_str)
    _set_run_font(run_data, size_pt=8, color_hex=MUTED_GRAY_HEX)

    # Navy line below header
    _add_horizontal_rule(doc, color_hex=NAVY_HEX, thickness=8)


def _add_footer(doc) -> None:
    """Footer: left 'ZYN Capital — Estruturação...' right 'Pág. X'."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Emu(9746 * 635))

    left_text = "ZYN Capital — Estruturação e Assessoria em Mercado de Capitais | Documento Confidencial — Circulação Restrita"
    run_left = p.add_run(left_text)
    _set_run_font(run_left, size_pt=6.5, color_hex=MUTED_GRAY_HEX, bold=False)

    run_tab = p.add_run("\t")
    _set_run_font(run_tab, size_pt=6.5, color_hex=MUTED_GRAY_HEX)

    run_pag = p.add_run("Pág. ")
    _set_run_font(run_pag, size_pt=6.5, color_hex=MUTED_GRAY_HEX)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run_field = p.add_run()
    _set_run_font(run_field, size_pt=6.5, color_hex=MUTED_GRAY_HEX)
    run_field._r.append(fld_char_begin)
    run_instr = p.add_run()
    _set_run_font(run_instr, size_pt=6.5, color_hex=MUTED_GRAY_HEX)
    run_instr._r.append(instr_text)
    run_end = p.add_run()
    _set_run_font(run_end, size_pt=6.5, color_hex=MUTED_GRAY_HEX)
    run_end._r.append(fld_char_end)


def _add_kpi_strip(doc, kpis: Dict[str, Any]) -> None:
    """5-column KPI strip with enhanced sizing and distinct column colors."""
    labels_keys = [
        ("PRINCIPAL", "principal"),
        ("PRAZO", "prazo"),
        ("TAXA ALL-IN", "taxa_all_in"),
        ("DSCR / LTV", "dscr_ltv"),
        ("RATING", "rating"),
    ]

    table = doc.add_table(rows=2, cols=5)
    _set_table_width(table)
    _remove_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_width = TABLE_WIDTH_DXA // 5

    for col_idx, (label, key) in enumerate(labels_keys):
        color_hex = KPI_COLORS[col_idx % len(KPI_COLORS)]
        value_str = kpis.get(key, "N/D")
        if isinstance(value_str, tuple):
            value_str = value_str[1]

        # Label row — taller padding
        cell_label = table.cell(0, col_idx)
        _set_cell_shading(cell_label, color_hex)
        _set_cell_no_borders(cell_label)
        _set_cell_margins(cell_label, top=100, left=100, bottom=20, right=80)
        _set_col_width(cell_label, col_width)
        cell_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _clear_cell(cell_label)
        p_label = cell_label.add_paragraph()
        p_label.paragraph_format.space_before = Pt(0)
        p_label.paragraph_format.space_after = Pt(0)
        run_label = p_label.add_run(label)
        _set_run_font(run_label, size_pt=7, color_hex=LIGHT_LABEL_HEX, bold=False)

        # Value row — larger text, more padding
        cell_val = table.cell(1, col_idx)
        _set_cell_shading(cell_val, color_hex)
        _set_cell_no_borders(cell_val)
        _set_cell_margins(cell_val, top=40, left=100, bottom=120, right=80)
        _set_col_width(cell_val, col_width)
        cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _clear_cell(cell_val)
        p_val = cell_val.add_paragraph()
        p_val.paragraph_format.space_before = Pt(0)
        p_val.paragraph_format.space_after = Pt(0)
        run_val = p_val.add_run(_safe(value_str))
        _set_run_font(run_val, size_pt=9, color_hex=WHITE_HEX, bold=True)

    _add_spacer(doc, 8)


def _add_section_heading(doc, title: str) -> None:
    """Full-width navy bar with 12pt bold white text — more prominent."""
    _add_spacer(doc, 12)
    table = doc.add_table(rows=1, cols=1)
    _set_table_width(table)
    _remove_table_borders(table)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, NAVY_HEX)
    _set_cell_margins(cell, top=160, left=220, bottom=160, right=220)
    _set_col_width(cell, TABLE_WIDTH_DXA)
    _clear_cell(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    _set_run_font(run, size_pt=12, color_hex=WHITE_HEX, bold=True)
    _add_spacer(doc, 6)


def _add_subsection_heading(doc, title: str) -> None:
    """Light gray bar with 9pt bold navy text."""
    _add_spacer(doc, 4)
    table = doc.add_table(rows=1, cols=1)
    _set_table_width(table)
    _remove_table_borders(table)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, LIGHT_GRAY_BG2_HEX)
    _set_cell_margins(cell, top=120, left=200, bottom=120, right=200)
    _set_col_width(cell, TABLE_WIDTH_DXA)
    _clear_cell(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    _set_run_font(run, size_pt=9, color_hex=NAVY_HEX, bold=True)
    _add_spacer(doc, 3)


def _add_critical_alert(doc, text: str) -> None:
    """Red banner for critical alerts."""
    table = doc.add_table(rows=1, cols=1)
    _set_table_width(table)
    _remove_table_borders(table)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, RED_LIGHT_BG_HEX)
    _set_cell_margins(cell, top=120, left=200, bottom=120, right=200)
    _set_col_width(cell, TABLE_WIDTH_DXA)
    _set_cell_borders(
        cell,
        left=("single", 24, RED_SOLID_HEX),
        top=("none", 0, RED_LIGHT_BG_HEX),
        bottom=("none", 0, RED_LIGHT_BG_HEX),
        right=("none", 0, RED_LIGHT_BG_HEX),
    )
    _clear_cell(cell)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("■ " + text)
    _set_run_font(run, size_pt=8.5, color_hex=DARK_RED_HEX, bold=True)


def _add_kv_table(doc, rows: List[Tuple[str, str]], label_width: int = 3600) -> None:
    """Key-value 2-column table — label column always #F2F4F6, thin borders, wider labels."""
    value_width = TABLE_WIDTH_DXA - label_width
    table = doc.add_table(rows=len(rows), cols=2)
    _set_table_width(table)
    _remove_table_borders(table)

    for i, (label, value) in enumerate(rows):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)

        _set_col_width(cell_label, label_width)
        _set_col_width(cell_value, value_width)

        # Label column always has fill; value column white
        _set_cell_shading(cell_label, LIGHT_GRAY_BG_HEX)
        _set_cell_shading(cell_value, WHITE_HEX)

        # Thin borders only
        border_spec = ("single", 2, BORDER_LIGHT_HEX)
        _set_cell_borders(cell_label, top=border_spec, bottom=border_spec,
                          left=border_spec, right=border_spec)
        _set_cell_borders(cell_value, top=border_spec, bottom=border_spec,
                          left=border_spec, right=border_spec)

        _set_cell_margins(cell_label, top=70, left=140, bottom=70, right=100)
        _set_cell_margins(cell_value, top=70, left=140, bottom=70, right=100)

        cell_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_value.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        _clear_cell(cell_label)
        p_l = cell_label.add_paragraph()
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after = Pt(0)
        run_l = p_l.add_run(_safe(label))
        _set_run_font(run_l, size_pt=8.5, color_hex=DARK_SLATE_HEX, bold=True)

        _clear_cell(cell_value)
        p_v = cell_value.add_paragraph()
        p_v.paragraph_format.space_before = Pt(0)
        p_v.paragraph_format.space_after = Pt(0)

        is_total = label.upper().startswith("TOTAL")
        is_bold = is_total or label.upper().startswith("AVAL")

        # TOTAL rows get special treatment
        if is_total:
            _set_cell_shading(cell_label, LIGHT_GRAY_BG2_HEX)
            _set_cell_shading(cell_value, LIGHT_GRAY_BG2_HEX)
            _set_cell_borders(cell_label, top=("single", 8, NAVY_HEX),
                              bottom=border_spec, left=border_spec, right=border_spec)
            _set_cell_borders(cell_value, top=("single", 8, NAVY_HEX),
                              bottom=border_spec, left=border_spec, right=border_spec)

        run_v = p_v.add_run(_safe(value))
        _set_run_font(run_v, size_pt=8.5, color_hex=NAVY_HEX, bold=is_bold)


def _add_data_table(
    doc,
    headers: List[str],
    rows: List[List[str]],
    col_widths: Optional[List[int]] = None,
    bold_last_row: bool = False,
    status_col: Optional[int] = None,
) -> None:
    """Data table: navy header (no borders, clean), alternating white/#F7F8FA rows, thin bottom border."""
    if not headers:
        return

    n_cols = len(headers)
    if col_widths is None:
        col_widths = [TABLE_WIDTH_DXA // n_cols] * n_cols

    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    _set_table_width(table)
    _remove_table_borders(table)

    # Header row — navy bg, white text, NO borders
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        _set_cell_shading(cell, NAVY_HEX)
        _set_cell_margins(cell, top=80, left=100, bottom=80, right=100)
        _set_col_width(cell, col_widths[col_idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_no_borders(cell)
        _clear_cell(cell)
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header_text)
        _set_run_font(run, size_pt=8.5, color_hex=WHITE_HEX, bold=True)

    # Data rows — alternating white / #F7F8FA, thin bottom border only
    for row_idx, row_data in enumerate(rows):
        is_last = row_idx == len(rows) - 1
        is_total = is_last and bold_last_row
        fill = ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX
        if is_total:
            fill = LIGHT_GRAY_BG2_HEX

        for col_idx in range(n_cols):
            cell = table.cell(row_idx + 1, col_idx)
            _set_cell_margins(cell, top=60, left=100, bottom=60, right=100)
            _set_col_width(cell, col_widths[col_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Thin bottom border only for clean look
            bottom_border = ("single", 4, BORDER_LIGHT_HEX)
            if is_total:
                # TOTAL row: navy top border
                _set_cell_borders(cell,
                                  top=("single", 8, NAVY_HEX),
                                  bottom=bottom_border,
                                  left=("none", 0, WHITE_HEX),
                                  right=("none", 0, WHITE_HEX))
            else:
                _set_cell_borders(cell,
                                  top=("none", 0, WHITE_HEX),
                                  bottom=bottom_border,
                                  left=("none", 0, WHITE_HEX),
                                  right=("none", 0, WHITE_HEX))

            text = _safe(row_data[col_idx] if col_idx < len(row_data) else "")
            _clear_cell(cell)
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # Status column with colored badge
            if status_col is not None and col_idx == status_col:
                text_upper = text.upper().strip()
                text_hex, badge_fill = STATUS_STYLES.get(
                    text_upper, (NAVY_HEX, fill)
                )
                if text in ("✅", "☑", "☑️"):
                    text_hex, badge_fill = WHITE_HEX, GREEN_SOLID_HEX
                    text = "✓"
                elif text_upper in ("SAUDÁVEL", "COBERTO", "EXCELENTE"):
                    text = "✓ " + text
                elif text_upper in ("ATENÇÃO", "MONITORAR"):
                    text = "⚠ " + text
                elif text_upper == "CRÍTICO":
                    text = "✕ " + text
                _set_cell_shading(cell, badge_fill)
                run = p.add_run(text)
                _set_run_font(run, size_pt=8.5, color_hex=text_hex, bold=True)
            else:
                _set_cell_shading(cell, fill)
                run = p.add_run(text)
                _set_run_font(run, size_pt=8.5, color_hex=NAVY_HEX, bold=is_total)


def _add_alert_box(doc, title: str, body: str, alert_type: str = "warning") -> None:
    """Alert box with thick left border (24pt), solid fill, more padding."""
    fill_hex, border_hex, title_hex = ALERT_STYLES.get(alert_type, ALERT_STYLES["warning"])

    _add_spacer(doc, 3)
    table = doc.add_table(rows=1, cols=1)
    _set_table_width(table)
    _remove_table_borders(table)

    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill_hex)
    _set_cell_margins(cell, top=140, left=240, bottom=140, right=220)
    _set_col_width(cell, TABLE_WIDTH_DXA)

    _set_cell_borders(
        cell,
        left=("single", 24, border_hex),
        top=("none", 0, fill_hex),
        bottom=("none", 0, fill_hex),
        right=("none", 0, fill_hex),
    )

    _clear_cell(cell)

    p_title = cell.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(3)
    run_title = p_title.add_run(title)
    _set_run_font(run_title, size_pt=9, color_hex=title_hex, bold=True)

    if body:
        p_body = cell.add_paragraph()
        p_body.paragraph_format.space_before = Pt(0)
        p_body.paragraph_format.space_after = Pt(0)
        run_body = p_body.add_run(_safe(body))
        _set_run_font(run_body, size_pt=8.5, color_hex=NAVY_HEX)
    _add_spacer(doc, 3)


def _add_body_text(doc, text: str) -> None:
    """Body paragraph — 9pt for better readability."""
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, size_pt=9, color_hex=NAVY_HEX)


def _add_risk_table(doc, risks: List[Dict[str, str]]) -> None:
    """Risk matrix with color-coded severity badges (solid fills, white text for high)."""
    headers = ["#", "RISCO", "PROB.", "IMPACTO", "MITIGANTE", "SEV."]
    col_widths = [380, 2100, 800, 800, 4066, 1600]

    n_rows = len(risks)
    table = doc.add_table(rows=1 + n_rows, cols=6)
    _set_table_width(table)
    _remove_table_borders(table)

    # Header
    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        _set_cell_shading(cell, NAVY_HEX)
        _set_cell_margins(cell, top=80, left=100, bottom=80, right=100)
        _set_col_width(cell, col_widths[col_idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_no_borders(cell)
        _clear_cell(cell)
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        _set_run_font(run, size_pt=8.5, color_hex=WHITE_HEX, bold=True)

    # Data rows
    for row_idx, risk in enumerate(risks):
        fill = ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX
        nivel = _safe(risk.get("nivel", risk.get("nível", risk.get("severidade", ""))), "N/D").upper()

        row_data = [
            str(row_idx + 1),
            _safe(risk.get("risco", risk.get("descricao", "")), "N/D"),
            _safe(risk.get("probabilidade", risk.get("prob", "")), "N/D"),
            _safe(risk.get("impacto", "")),
            _safe(risk.get("mitigante", "")),
            nivel,
        ]

        for col_idx in range(6):
            cell = table.cell(row_idx + 1, col_idx)
            _set_cell_margins(cell, top=60, left=100, bottom=60, right=100)
            _set_col_width(cell, col_widths[col_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _clear_cell(cell)

            # Bottom border only
            _set_cell_borders(cell,
                              top=("none", 0, WHITE_HEX),
                              bottom=("single", 4, BORDER_LIGHT_HEX),
                              left=("none", 0, WHITE_HEX),
                              right=("none", 0, WHITE_HEX))

            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            text = row_data[col_idx]

            if col_idx == 5:  # Severity badge — solid colors
                matched = False
                for sev_key, (sev_text_hex, sev_fill) in RISK_LEVEL_STYLES.items():
                    if sev_key in nivel:
                        _set_cell_shading(cell, sev_fill)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(text)
                        _set_run_font(run, size_pt=8, color_hex=sev_text_hex, bold=True)
                        matched = True
                        break
                if not matched:
                    _set_cell_shading(cell, fill)
                    run = p.add_run(text)
                    _set_run_font(run, size_pt=8, color_hex=NAVY_HEX, bold=True)
            else:
                _set_cell_shading(cell, fill)
                run = p.add_run(text)
                _set_run_font(run, size_pt=8.5 if col_idx > 0 else 8, color_hex=NAVY_HEX,
                              bold=(col_idx == 0))


def _add_covenant_table(doc, covenants: List[Dict[str, str]]) -> None:
    """Covenant table: monitoramento column right-aligned in gray italic."""
    if not covenants:
        return
    headers = ["COVENANT", "PARÂMETRO", "MONITORAMENTO"]
    col_widths = [2800, 4446, 2500]

    table = doc.add_table(rows=1 + len(covenants), cols=3)
    _set_table_width(table)
    _remove_table_borders(table)

    for col_idx, h in enumerate(headers):
        cell = table.cell(0, col_idx)
        _set_cell_shading(cell, NAVY_HEX)
        _set_cell_margins(cell, top=80, left=100, bottom=80, right=100)
        _set_col_width(cell, col_widths[col_idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_no_borders(cell)
        _clear_cell(cell)
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        _set_run_font(run, size_pt=8.5, color_hex=WHITE_HEX, bold=True)

    for row_idx, cov in enumerate(covenants):
        fill = ALT_ROW_HEX if row_idx % 2 == 0 else WHITE_HEX
        row_data = [
            _safe(cov.get("covenant", cov.get("nome", ""))),
            _safe(cov.get("parametro", cov.get("parâmetro", ""))),
            _safe(cov.get("monitoramento", cov.get("frequencia", ""))),
        ]
        for col_idx in range(3):
            cell = table.cell(row_idx + 1, col_idx)
            _set_cell_shading(cell, fill)
            _set_cell_margins(cell, top=60, left=100, bottom=60, right=100)
            _set_col_width(cell, col_widths[col_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_borders(cell,
                              top=("none", 0, WHITE_HEX),
                              bottom=("single", 4, BORDER_LIGHT_HEX),
                              left=("none", 0, WHITE_HEX),
                              right=("none", 0, WHITE_HEX))
            _clear_cell(cell)
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            text = row_data[col_idx]
            if col_idx == 2:
                # Monitoramento: right-aligned, gray italic
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(text)
                _set_run_font(run, size_pt=8.5, color_hex=MUTED_GRAY_HEX, italic=True)
            elif col_idx == 0:
                run = p.add_run(text)
                _set_run_font(run, size_pt=8.5, color_hex=NAVY_HEX, bold=True)
            else:
                run = p.add_run(text)
                _set_run_font(run, size_pt=8.5, color_hex=NAVY_HEX)


# ═══════════════════════════════════════════════════════════════════════════
# Section Renderers — MAC v4.1 (13 sections)
# ═══════════════════════════════════════════════════════════════════════════

def _render_section_tomador(doc, data: dict, parametros: dict) -> None:
    """I. TOMADOR — Ficha do tomador com dados extras."""
    _add_section_heading(doc, SECTION_TITLES["tomador"])

    ficha = [
        ("Razão Social", _safe(parametros.get("tomador"))),
        ("CNPJ", _safe(parametros.get("cnpj"))),
        ("Setor", _safe(parametros.get("setor", data.get("setor", "")))),
        ("Tipo Operação", _safe(parametros.get("tipo_operacao"))),
        ("Instrumento", _safe(parametros.get("instrumento", parametros.get("tipo_operacao")))),
        ("Sócio Responsável ZYN", _safe(parametros.get("socio_responsavel"))),
    ]

    extra_fields = data.get("dados", data.get("ficha", {}))
    if isinstance(extra_fields, dict):
        for k, v in extra_fields.items():
            if k not in ("analise", "rating_secao", "flags", "setor"):
                ficha.append((k, _safe(v)))

    _add_subsection_heading(doc, "Ficha do Tomador")
    _add_kv_table(doc, ficha)

    analise_text = data.get("analise", "")
    if analise_text:
        _add_alert_box(doc, "Análise", analise_text, "info")

    _render_flags(doc, data.get("flags", []))


def _render_section_patrimonio(doc, data: dict) -> None:
    """II. PATRIMÔNIO — Ativos, matrículas, valuations."""
    _add_section_heading(doc, SECTION_TITLES["patrimonio"])

    analise_text = data.get("analise", "")
    if analise_text:
        _add_alert_box(doc, "Análise Patrimonial", analise_text, "info")

    for sub_key in ("patrimoniopf", "patrimonio_pf", "pf"):
        sub = data.get(sub_key, {})
        if isinstance(sub, dict) and sub.get("analise"):
            _add_subsection_heading(doc, "Patrimônio Pessoa Física")
            _add_body_text(doc, sub["analise"])

    for sub_key in ("patrimoniocorp", "patrimonio_corp", "corporativo"):
        sub = data.get(sub_key, {})
        if isinstance(sub, dict) and sub.get("analise"):
            _add_subsection_heading(doc, "Patrimônio Corporativo")
            _add_body_text(doc, sub["analise"])

    matriculas = data.get("matriculas", data.get("tabela_matriculas", []))
    if isinstance(matriculas, list) and matriculas:
        _add_subsection_heading(doc, "Matrículas e Avaliações")
        if isinstance(matriculas[0], dict):
            headers = list(matriculas[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in matriculas]
            _add_data_table(doc, headers, rows, bold_last_row=True)

    tabela = data.get("tabela", data.get("dados", []))
    if isinstance(tabela, list) and tabela and not matriculas:
        if isinstance(tabela[0], dict):
            headers = list(tabela[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in tabela]
            _add_data_table(doc, headers, rows)
        elif isinstance(tabela[0], list):
            _add_data_table(doc, tabela[0], tabela[1:])

    alerta = data.get("alerta_irpf", data.get("alerta", ""))
    if alerta:
        _add_alert_box(doc, "⚠ Alerta Patrimonial", alerta, "warning")

    _render_flags(doc, data.get("flags", []))


def _render_section_producao(doc, data: dict) -> None:
    """III. PRODUÇÃO — Faturamento, safras, VGV, VSO."""
    _add_section_heading(doc, SECTION_TITLES["producao"])

    analise_text = data.get("analise", "")
    if analise_text:
        _add_alert_box(doc, "Análise Operacional", analise_text, "info")

    tabela = data.get("tabela", data.get("tabela_safra", data.get("dados", [])))
    if isinstance(tabela, list) and tabela:
        sub_title = data.get("subtitulo", "Histórico de Produção")
        _add_subsection_heading(doc, sub_title)
        if isinstance(tabela[0], dict):
            headers = list(tabela[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in tabela]
            _add_data_table(doc, headers, rows, bold_last_row=True)

    projecoes = data.get("projecoes", data.get("tabela_projecoes", []))
    if isinstance(projecoes, list) and projecoes:
        _add_subsection_heading(doc, "Projeções de Produção")
        if isinstance(projecoes[0], dict):
            headers = list(projecoes[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in projecoes]
            _add_data_table(doc, headers, rows, bold_last_row=True)

    nota = data.get("nota", "")
    if nota:
        _add_alert_box(doc, "Nota", nota, "info")

    _render_flags(doc, data.get("flags", []))


def _render_section_capital(doc, data: dict) -> None:
    """IV. CAPITAL E ESTRUTURA DE ENDIVIDAMENTO."""
    _add_section_heading(doc, SECTION_TITLES["capital"])

    analise_text = data.get("analise", "")
    if analise_text:
        _add_alert_box(doc, "Análise de Capital", analise_text, "info")

    tabela = data.get("tabela", data.get("dados", []))
    if isinstance(tabela, list) and tabela:
        _add_subsection_heading(doc, "Estrutura de Endividamento")
        if isinstance(tabela[0], dict):
            headers = list(tabela[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in tabela]
            _add_data_table(doc, headers, rows, bold_last_row=True)

    _render_flags(doc, data.get("flags", []))


def _render_section_operacao(doc, data: dict, parametros: dict) -> None:
    """V. OPERAÇÃO — Parâmetros da operação."""
    _add_section_heading(doc, SECTION_TITLES["operacao"])

    ficha = [
        ("Instrumento", _safe(parametros.get("instrumento", parametros.get("tipo_operacao")))),
        ("Volume", _format_currency(parametros.get("volume"))),
        ("Prazo", f"{_safe(parametros.get('prazo_meses'))} meses"),
        ("Taxa", _safe(parametros.get("taxa"))),
        ("Amortização", _safe(parametros.get("amortizacao"))),
        ("Garantias", ", ".join(parametros.get("garantias", [])) or "N/D"),
    ]

    extra = data.get("dados", data.get("parametros_extra", {}))
    if isinstance(extra, dict):
        for k, v in extra.items():
            if k not in ("analise", "flags"):
                ficha.append((k, _safe(v)))

    _add_subsection_heading(doc, "Parâmetros da Operação")
    _add_kv_table(doc, ficha, label_width=3600)

    analise_text = data.get("analise", "")
    if analise_text:
        _add_alert_box(doc, "Análise da Operação", analise_text, "info")

    _render_flags(doc, data.get("flags", []))


def _render_section_estrutura(doc, data: dict) -> None:
    """VI. ESTRUTURA DA OPERAÇÃO — Detalhamento bilateral/emissão."""
    _add_section_heading(doc, SECTION_TITLES["estrutura"])

    subtitulo = data.get("subtitulo", "")
    if subtitulo:
        _add_subsection_heading(doc, subtitulo)

    analise_text = data.get("analise", "")
    if analise_text:
        _add_body_text(doc, analise_text)

    ficha = data.get("ficha", data.get("dados", []))
    if isinstance(ficha, list) and ficha:
        if isinstance(ficha[0], (list, tuple)):
            _add_kv_table(doc, ficha, label_width=3600)
        elif isinstance(ficha[0], dict):
            headers = list(ficha[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in ficha]
            _add_data_table(doc, headers, rows)
    elif isinstance(ficha, dict):
        kv_rows = [(k, _safe(v)) for k, v in ficha.items() if k not in ("analise", "flags")]
        if kv_rows:
            _add_kv_table(doc, kv_rows, label_width=3600)

    fases = data.get("fases", [])
    if isinstance(fases, list) and fases:
        _add_subsection_heading(doc, "Fases da Operação")
        if isinstance(fases[0], dict):
            headers = list(fases[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in fases]
            _add_data_table(doc, headers, rows)

    _render_flags(doc, data.get("flags", []))


def _render_section_pagamento(doc, data: dict) -> None:
    """VII. MODELAGEM FINANCEIRA — CRONOGRAMA."""
    _add_section_heading(doc, SECTION_TITLES["pagamento"])

    analise_text = data.get("analise", "")
    if analise_text:
        _add_body_text(doc, analise_text)

    cronograma = data.get("cronograma", data.get("tabela", data.get("dados", [])))
    if isinstance(cronograma, list) and cronograma:
        _add_subsection_heading(doc, "Fluxo de Pagamentos")
        if isinstance(cronograma[0], dict):
            headers = list(cronograma[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in cronograma]
            _add_data_table(doc, headers, rows, bold_last_row=True)
        elif isinstance(cronograma[0], list):
            _add_data_table(doc, cronograma[0], cronograma[1:], bold_last_row=True)

    totais = data.get("totais", data.get("resumo_totais", []))
    if isinstance(totais, list) and totais:
        _add_subsection_heading(doc, "Resumo de Custos")
        if isinstance(totais[0], dict):
            headers = list(totais[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in totais]
            _add_data_table(doc, headers, rows)

    destinacao = data.get("destinacao", data.get("uso_recursos", []))
    if isinstance(destinacao, list) and destinacao:
        _add_subsection_heading(doc, "Destinação dos Recursos")
        if isinstance(destinacao[0], dict):
            headers = list(destinacao[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in destinacao]
            _add_data_table(doc, headers, rows, bold_last_row=True)

    _render_flags(doc, data.get("flags", []))


def _render_section_capacidade(doc, data: dict) -> None:
    """VIII. ANÁLISE FINANCEIRA — CAPACIDADE DE PAGAMENTO."""
    _add_section_heading(doc, SECTION_TITLES["capacidade"])

    analise_text = data.get("analise", "")
    if analise_text:
        _add_alert_box(doc, "Capacidade de Pagamento", analise_text, "positive")

    cobertura = data.get("cobertura", data.get("tabela", []))
    if isinstance(cobertura, list) and cobertura:
        _add_subsection_heading(doc, "Cobertura Anual")
        if isinstance(cobertura[0], dict):
            headers = list(cobertura[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in cobertura]
            status_idx = None
            for i, h in enumerate(headers):
                if h.lower() in ("status", "sit."):
                    status_idx = i
                    break
            _add_data_table(doc, headers, rows, status_col=status_idx)

    sensibilidade = data.get("sensibilidade", data.get("sensibilidade_receita", []))
    if isinstance(sensibilidade, list) and sensibilidade:
        _add_subsection_heading(doc, "Análise de Sensibilidade — Variação de Receita")
        if isinstance(sensibilidade[0], dict):
            headers = list(sensibilidade[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in sensibilidade]
            status_idx = None
            for i, h in enumerate(headers):
                if h.lower() in ("status", "sit."):
                    status_idx = i
                    break
            _add_data_table(doc, headers, rows, status_col=status_idx)

    sens_cdi = data.get("sensibilidade_cdi", [])
    if isinstance(sens_cdi, list) and sens_cdi:
        _add_subsection_heading(doc, "Sensibilidade ao CDI")
        if isinstance(sens_cdi[0], dict):
            headers = list(sens_cdi[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in sens_cdi]
            _add_data_table(doc, headers, rows)

    nota = data.get("nota", "")
    if nota:
        _add_alert_box(doc, "Nota sobre sensibilidade", nota, "info")

    _render_flags(doc, data.get("flags", []))


def _render_section_impacto(doc, data: dict) -> None:
    """IX. IMPACTO FINANCEIRO — Before/after comparison."""
    _add_section_heading(doc, SECTION_TITLES["impacto"])

    analise_text = data.get("analise", "")
    if analise_text:
        _add_body_text(doc, analise_text)

    tabela = data.get("tabela", data.get("dados", []))
    if isinstance(tabela, list) and tabela:
        _add_subsection_heading(doc, "Impacto Financeiro Imediato")
        if isinstance(tabela[0], dict):
            headers = list(tabela[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in tabela]
            status_idx = None
            for i, h in enumerate(headers):
                if h.lower() in ("status", "sit."):
                    status_idx = i
                    break
            _add_data_table(doc, headers, rows, status_col=status_idx)

    _render_flags(doc, data.get("flags", []))


def _render_section_onus(doc, data: dict) -> None:
    """X. MAPA DE ÔNUS E GARANTIAS."""
    _add_section_heading(doc, SECTION_TITLES["onus"])

    alerta_critico = data.get("alerta_critico", "")
    if alerta_critico:
        _add_critical_alert(doc, alerta_critico)

    analise_text = data.get("analise", "")
    if analise_text:
        _add_body_text(doc, analise_text)

    tabela = data.get("tabela", data.get("onus", data.get("dados", [])))
    if isinstance(tabela, list) and tabela:
        if isinstance(tabela[0], dict):
            headers = list(tabela[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in tabela]
            _add_data_table(doc, headers, rows, bold_last_row=True)

    garantias = data.get("garantias", [])
    if isinstance(garantias, list) and garantias and not tabela:
        if isinstance(garantias[0], dict):
            headers = list(garantias[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in garantias]
            _add_data_table(doc, headers, rows)

    alerta_amb = data.get("alerta_ambiental", "")
    if alerta_amb:
        _add_alert_box(doc, "⚠ Alerta Ambiental — Condição Suspensiva", alerta_amb, "warning")

    _render_flags(doc, data.get("flags", []))


def _render_section_riscos(doc, data: dict) -> None:
    """XI. RISCOS, MITIGAÇÕES E COVENANTS."""
    _add_section_heading(doc, SECTION_TITLES["riscos"])

    analise_text = data.get("analise", "")
    if analise_text:
        _add_alert_box(doc, "Análise de Riscos", analise_text, "risk")

    riscos = data.get("riscos", data.get("matriz", data.get("tabela", [])))
    if isinstance(riscos, list) and riscos:
        _add_subsection_heading(doc, "Matriz de Riscos")
        _add_risk_table(doc, riscos)

    covenants = data.get("covenants", [])
    if isinstance(covenants, list) and covenants:
        _add_subsection_heading(doc, "Covenants e Condicionantes")
        _add_covenant_table(doc, covenants)

    _render_flags(doc, data.get("flags", []))


def _render_section_cronograma(doc, data: dict) -> None:
    """XII. CRONOGRAMA DE EXECUÇÃO."""
    _add_section_heading(doc, SECTION_TITLES["cronograma"])

    analise_text = data.get("analise", "")
    if analise_text:
        _add_body_text(doc, analise_text)

    tabela = data.get("tabela", data.get("dados", data.get("fases", [])))
    if isinstance(tabela, list) and tabela:
        if isinstance(tabela[0], dict):
            headers = list(tabela[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in tabela]
            _add_data_table(doc, headers, rows)

    _render_flags(doc, data.get("flags", []))


def _render_section_conclusao(doc, data: dict, parametros: dict) -> None:
    """XIII. CONCLUSÃO E RECOMENDAÇÃO — ZYN CAPITAL."""
    _add_section_heading(doc, SECTION_TITLES["conclusao"])

    status_text = data.get("status", "")
    if status_text:
        _add_subsection_heading(doc, f"— {status_text}")

    analise_text = data.get("analise", data.get("conclusao", ""))
    if analise_text:
        _add_body_text(doc, analise_text)

    recomendacao = data.get("recomendacao", "")
    if recomendacao:
        _add_alert_box(doc, "Recomendação", recomendacao, "positive")

    _render_flags(doc, data.get("flags", []))


def _render_section_investidores(doc, data: dict) -> None:
    """Investor matching section — rendered if available."""
    if not data:
        return

    _add_section_heading(doc, "INVESTOR MATCHING — INVESTIDORES SUGERIDOS")

    investidores = data.get("investidores", data.get("tabela", []))
    if isinstance(investidores, list) and investidores:
        if isinstance(investidores[0], dict):
            headers = list(investidores[0].keys())
            rows = [[_safe(item.get(h)) for h in headers] for item in investidores]
            _add_data_table(doc, headers, rows)

    nota = data.get("nota", "")
    if nota:
        _add_body_text(doc, nota)


def _render_flags(doc, flags) -> None:
    """Consolidate all flags into ONE alert box with bullet points."""
    if not flags:
        return
    if isinstance(flags, str):
        flags = [flags]
    if not flags:
        return

    if len(flags) == 1:
        _add_alert_box(doc, "⚠ Flag", flags[0], "warning")
    else:
        # Consolidate into single box with bullets
        body_lines = []
        for f in flags:
            body_lines.append(f"• {f}")
        body = "\n".join(body_lines)
        _add_alert_box(doc, f"⚠ Flags ({len(flags)})", body, "warning")


def _render_rating_final(doc, analise: dict) -> None:
    """Render the final rating as a prominent badge box."""
    rating = analise.get("rating_final", {})
    if not rating:
        return

    nota = _safe(rating.get("nota", rating.get("rating", "")), "N/D")
    descricao = _safe(rating.get("descricao", rating.get("analise", "")), "")

    # Create a 2-column table: large rating badge | description
    _add_spacer(doc, 8)
    _add_subsection_heading(doc, "RATING FINAL")

    table = doc.add_table(rows=1, cols=2)
    _set_table_width(table)
    _remove_table_borders(table)

    badge_width = 1800
    desc_width = TABLE_WIDTH_DXA - badge_width

    # Rating badge cell — large letter with colored background
    text_hex, fill_hex = RATING_BADGE_COLORS.get(nota, (WHITE_HEX, NAVY_HEX))
    cell_badge = table.cell(0, 0)
    _set_cell_shading(cell_badge, fill_hex)
    _set_cell_margins(cell_badge, top=200, left=100, bottom=200, right=100)
    _set_col_width(cell_badge, badge_width)
    cell_badge.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_no_borders(cell_badge)
    _clear_cell(cell_badge)
    p_badge = cell_badge.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_before = Pt(0)
    p_badge.paragraph_format.space_after = Pt(0)
    run_badge = p_badge.add_run(nota)
    _set_run_font(run_badge, size_pt=28, color_hex=text_hex, bold=True)

    # Description cell
    cell_desc = table.cell(0, 1)
    _set_cell_shading(cell_desc, LIGHT_GRAY_BG_HEX)
    _set_cell_margins(cell_desc, top=160, left=200, bottom=160, right=200)
    _set_col_width(cell_desc, desc_width)
    cell_desc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_no_borders(cell_desc)
    _clear_cell(cell_desc)

    p_label = cell_desc.add_paragraph()
    p_label.paragraph_format.space_before = Pt(0)
    p_label.paragraph_format.space_after = Pt(4)
    run_label = p_label.add_run(f"RATING: {nota}")
    _set_run_font(run_label, size_pt=11, color_hex=NAVY_HEX, bold=True)

    if descricao:
        p_desc = cell_desc.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(0)
        p_desc.paragraph_format.space_after = Pt(0)
        run_desc = p_desc.add_run(descricao)
        _set_run_font(run_desc, size_pt=8.5, color_hex=DARK_SLATE_HEX)

    _add_spacer(doc, 6)


# ═══════════════════════════════════════════════════════════════════════════
# Main Function — MAC v4.1
# ═══════════════════════════════════════════════════════════════════════════

def generate_mac(analise: dict, parametros: dict, output_path: str) -> str:
    """
    Generate a MAC (Memorando de Análise de Crédito) .docx file — v4.1 enhanced.

    Args:
        analise: Analysis JSON with section keys (tomador, patrimonio, producao,
                 capital, operacao, estrutura, pagamento, capacidade, impacto,
                 onus, riscos, cronograma, conclusao, investor_matching,
                 rating_final, kpis).
        parametros: Operation parameters dict.
        output_path: Path where the .docx will be saved.

    Returns:
        The output_path of the generated file.
    """
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.59)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Empty header
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        header.paragraphs[0].clear()

    # Footer
    _add_footer(doc)

    # ── 1. Branded header bar ──────────────────────────────────────────
    _add_header_bar(doc, parametros)

    # ── 2. Title — 18pt bold navy ─────────────────────────────────────
    _add_spacer(doc, 8)
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(2)
    p_title.paragraph_format.space_before = Pt(0)
    run_title = p_title.add_run("MEMORANDO DE ANÁLISE DE CRÉDITO")
    _set_run_font(run_title, size_pt=18, color_hex=NAVY_HEX, bold=True)

    # 2pt navy underline below title
    _add_horizontal_rule(doc, color_hex=NAVY_HEX, thickness=6)

    # ── 3. Subtitle — 12pt dark slate ─────────────────────────────────
    tipo_op = _safe(parametros.get("tipo_operacao"), "")
    tomador = _safe(parametros.get("tomador"), "N/D")

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(2)
    p_sub.paragraph_format.space_before = Pt(4)
    sub_text = f"{tipo_op} — {tomador}" if tipo_op else tomador
    run_sub = p_sub.add_run(sub_text)
    _set_run_font(run_sub, size_pt=12, color_hex=DARK_SLATE_HEX)

    # ── 4. Meta-line ──────────────────────────────────────────────────
    taxa = _safe(parametros.get("taxa"), "N/D")
    prazo = f"{_safe(parametros.get('prazo_meses'), 'N/D')}m"
    volume_str = _format_currency(parametros.get("volume"))
    data_str = datetime.now().strftime("%d/%m/%Y")

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(10)
    p_meta.paragraph_format.space_before = Pt(0)

    meta_prefix = f"{volume_str} · {taxa} · {prazo} · "
    run_meta = p_meta.add_run(meta_prefix)
    _set_run_font(run_meta, size_pt=8.5, color_hex=MUTED_GRAY_HEX)
    run_date = p_meta.add_run(data_str)
    _set_run_font(run_date, size_pt=8.5, color_hex=MUTED_GRAY_HEX, italic=True)

    # ── 5. KPI Strip ─────────────────────────────────────────────────
    kpis_data = analise.get("kpis", {})
    if not kpis_data:
        volume = parametros.get("volume")
        prazo_m = parametros.get("prazo_meses")
        kpis_data = {
            "principal": _format_currency(volume) if volume else "N/D",
            "prazo": f"{prazo_m} meses" if prazo_m else "N/D",
            "taxa_all_in": _safe(parametros.get("taxa")),
            "dscr_ltv": "N/D",
            "rating": "N/D",
        }

    # Inject rating into KPIs if available
    rating_final = analise.get("rating_final", {})
    if rating_final and kpis_data.get("rating", "N/D") == "N/D":
        kpis_data["rating"] = _safe(rating_final.get("nota", rating_final.get("rating", "")))

    _add_kpi_strip(doc, kpis_data)

    # ── 6. Summary tables (optional — 3-col summary like Baron55) ────
    resumo = analise.get("resumo", {})
    if isinstance(resumo, dict):
        for sub_key, sub_title in [
            ("tomador_resumo", "Resumo do Tomador"),
            ("instrumento_resumo", "Resumo do Instrumento"),
            ("kpis_resumo", "KPIs da Operação"),
        ]:
            sub = resumo.get(sub_key, [])
            if isinstance(sub, list) and sub:
                _add_subsection_heading(doc, sub_title)
                if isinstance(sub[0], (list, tuple)):
                    _add_kv_table(doc, sub)

    # ── 7. Sections I-XIII ───────────────────────────────────────────
    section_renderers = {
        "tomador": lambda d: _render_section_tomador(doc, d, parametros),
        "patrimonio": lambda d: _render_section_patrimonio(doc, d),
        "producao": lambda d: _render_section_producao(doc, d),
        "capital": lambda d: _render_section_capital(doc, d),
        "operacao": lambda d: _render_section_operacao(doc, d, parametros),
        "estrutura": lambda d: _render_section_estrutura(doc, d),
        "pagamento": lambda d: _render_section_pagamento(doc, d),
        "capacidade": lambda d: _render_section_capacidade(doc, d),
        "impacto": lambda d: _render_section_impacto(doc, d),
        "onus": lambda d: _render_section_onus(doc, d),
        "riscos": lambda d: _render_section_riscos(doc, d),
        "cronograma": lambda d: _render_section_cronograma(doc, d),
        "conclusao": lambda d: _render_section_conclusao(doc, d, parametros),
    }

    for key in SECTION_TITLES:
        section_data = analise.get(key, {})
        if isinstance(section_data, dict) and section_data:
            renderer = section_renderers.get(key)
            if renderer:
                renderer(section_data)
            # Add thin separator line between major sections
            _add_horizontal_rule(doc, color_hex=SEPARATOR_HEX, thickness=4)

    # ── Backward compat: old keys ────────────────────────────────────
    if "covenants" in analise and "riscos" not in analise:
        cov_data = analise["covenants"]
        if isinstance(cov_data, dict):
            _render_section_riscos(doc, {"covenants": cov_data.get("tabela", cov_data.get("dados", []))})

    # ── Investor Matching ────────────────────────────────────────────
    investor_data = analise.get("investor_matching", {})
    if investor_data:
        _render_section_investidores(doc, investor_data)

    # ── Rating Final — Prominent Badge ────────────────────────────────
    _render_rating_final(doc, analise)

    # ── Closing — Author Block ────────────────────────────────────────
    _add_spacer(doc, 16)
    _add_horizontal_rule(doc, color_hex=BORDER_GRAY_HEX, thickness=4)
    _add_spacer(doc, 8)

    # "Trabalho elaborado:" label in gray
    p_elab = doc.add_paragraph()
    p_elab.paragraph_format.space_before = Pt(0)
    p_elab.paragraph_format.space_after = Pt(4)
    run_elab = p_elab.add_run("Trabalho elaborado:")
    _set_run_font(run_elab, size_pt=8, color_hex=MUTED_GRAY_HEX, italic=True)

    # Author name — bold navy
    socio = _safe(parametros.get("socio_responsavel"), "Danilo Salasar")
    p_author = doc.add_paragraph()
    p_author.paragraph_format.space_before = Pt(0)
    p_author.paragraph_format.space_after = Pt(2)
    run_nome = p_author.add_run(socio)
    _set_run_font(run_nome, size_pt=10, color_hex=NAVY_HEX, bold=True)

    # Title
    p_title_author = doc.add_paragraph()
    p_title_author.paragraph_format.space_before = Pt(0)
    p_title_author.paragraph_format.space_after = Pt(2)
    run_title_a = p_title_author.add_run("Sócio — ZYN Capital")
    _set_run_font(run_title_a, size_pt=8.5, color_hex=DARK_SLATE_HEX)

    # Email as blue link
    p_email = doc.add_paragraph()
    p_email.paragraph_format.space_before = Pt(0)
    p_email.paragraph_format.space_after = Pt(1)
    run_email = p_email.add_run("danilo@zyncapital.com.br")
    _set_run_font(run_email, size_pt=8, color_hex=DARK_BLUE_HEX)
    run_email.underline = True

    # Phone
    p_phone = doc.add_paragraph()
    p_phone.paragraph_format.space_before = Pt(0)
    p_phone.paragraph_format.space_after = Pt(0)
    run_phone = p_phone.add_run("65 9 9987-8781")
    _set_run_font(run_phone, size_pt=8, color_hex=MUTED_GRAY_HEX)

    # ZYN Capital centered block with tagline
    _add_spacer(doc, 16)
    _add_horizontal_rule(doc, color_hex=NAVY_HEX, thickness=6)
    _add_spacer(doc, 8)

    p_close = doc.add_paragraph()
    p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_close.paragraph_format.space_before = Pt(0)
    p_close.paragraph_format.space_after = Pt(4)
    run_zyn = p_close.add_run("ZYN Capital")
    _set_run_font(run_zyn, size_pt=14, color_hex=NAVY_HEX, bold=True)

    p_tagline = doc.add_paragraph()
    p_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tagline.paragraph_format.space_before = Pt(0)
    p_tagline.paragraph_format.space_after = Pt(2)
    run_tag = p_tagline.add_run("Estruturação e Assessoria em Mercado de Capitais")
    _set_run_font(run_tag, size_pt=9, color_hex=MUTED_GRAY_HEX)

    p_city = doc.add_paragraph()
    p_city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_city.paragraph_format.space_before = Pt(0)
    p_city.paragraph_format.space_after = Pt(4)
    run_city = p_city.add_run(f"São Paulo, {datetime.now().strftime('%B de %Y').lower()}")
    _set_run_font(run_city, size_pt=8, color_hex=MUTED_GRAY_HEX, italic=True)

    # Sources
    p_sources = doc.add_paragraph()
    p_sources.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sources.paragraph_format.space_before = Pt(8)
    p_sources.paragraph_format.space_after = Pt(4)
    run_src = p_sources.add_run(
        "Fontes: Documentos fornecidos pelo grupo | Demonstrações financeiras | "
        "Modelagem ZYN Capital"
    )
    _set_run_font(run_src, size_pt=6.5, color_hex=MUTED_GRAY_HEX, italic=True)

    # ── Save ─────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path
